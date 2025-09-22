import os
import io
import base64
import pandas as pd
import streamlit as st
from sqlalchemy import create_engine, text, inspect
from docx import Document
import streamlit.components.v1 as components
import requests

st.set_page_config(page_title="Data Importer", layout="wide")
st.title("Data Importer")

db_host = os.environ.get("DB_HOST", "localhost")
db_port = os.environ.get("DB_PORT", "3306")
db_name = os.environ.get("DB_NAME", "appdb")
db_user = os.environ.get("DB_USER", "appuser")
db_password = os.environ.get("DB_PASSWORD", "apppassword")
rag_base = os.environ.get("RAG_BASE_URL", "http://localhost:8001")

engine = create_engine(f"mysql+pymysql://{db_user}:{db_password}@{db_host}:{db_port}/{db_name}?charset=utf8mb4", pool_pre_ping=True)

def ensure_files_table():
    with engine.begin() as conn:
        conn.execute(text("CREATE TABLE IF NOT EXISTS files (id INT AUTO_INCREMENT PRIMARY KEY,filename VARCHAR(255) NOT NULL,content_type VARCHAR(128) NOT NULL,size_bytes BIGINT NOT NULL,data LONGBLOB NOT NULL,created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"))

def list_tables():
    insp = inspect(engine)
    return sorted([t for t in insp.get_table_names()])

def normalize_table_name(name):
    s = "".join(c if c.isalnum() or c in ("_",) else "_" for c in name.lower())
    s = s.strip("_")
    if s and s[0].isdigit():
        s = f"t_{s}"
    return s or "table"

def unique_table_name(base):
    base = normalize_table_name(base)
    existing = set(list_tables())
    if base not in existing:
        return base
    i = 2
    while f"{base}_{i}" in existing:
        i += 1
    return f"{base}_{i}"

def read_tabular_file(uploaded):
    name = uploaded.name.lower()
    data = uploaded.read()
    if name.endswith(".csv"):
        return pd.read_csv(io.BytesIO(data))
    if name.endswith(".xlsx") or name.endswith(".xls"):
        return pd.read_excel(io.BytesIO(data))
    return None

def write_df(df, table_name, replace=False):
    if replace and table_name in list_tables():
        with engine.begin() as conn:
            conn.execute(text(f"DROP TABLE IF EXISTS `{table_name}`"))
    df.to_sql(table_name, engine, index=False, if_exists="fail")

def save_file_to_db(uploaded):
    ensure_files_table()
    b = uploaded.read()
    with engine.begin() as conn:
        conn.execute(text("INSERT INTO files (filename, content_type, size_bytes, data) VALUES (:f,:c,:s,:d)"), {"f": uploaded.name, "c": uploaded.type or "", "s": len(b), "d": b})

def list_files():
    ensure_files_table()
    with engine.begin() as conn:
        rows = conn.execute(text("SELECT id, filename, content_type, size_bytes, created_at FROM files ORDER BY id DESC")).fetchall()
    return rows

def get_file_blob(file_id):
    with engine.begin() as conn:
        row = conn.execute(text("SELECT id, filename, content_type, size_bytes, data FROM files WHERE id=:i"), {"i": file_id}).fetchone()
    return row

def rag_ingest_files():
    try:
        r = requests.post(f"{rag_base}/ingest/files", timeout=180)
        return True, r.json()
    except Exception as e:
        return False, str(e)

def rag_ingest_tables(tables_csv=None):
    try:
        payload = {"tables": tables_csv} if tables_csv else {}
        r = requests.post(f"{rag_base}/ingest/db", json=payload, timeout=180)
        return True, r.json()
    except Exception as e:
        return False, str(e)

if "auto_sync" not in st.session_state:
    st.session_state.auto_sync = True

st.sidebar.checkbox("Auto-sync to RAG on upload/import", value=st.session_state.auto_sync, key="auto_sync")

tabs = st.tabs(["Upload Data", "Browse Tables", "Upload Files", "Files", "Clear All", "Chat"])

with tabs[0]:
    uploaded = st.file_uploader("Upload CSV/XLS/XLSX", type=["csv", "xls", "xlsx"])
    if uploaded:
        df = read_tabular_file(uploaded)
        if df is None:
            st.error("Unsupported file")
        else:
            st.subheader("Preview")
            st.dataframe(df.head(100), use_container_width=True)
            default_name = os.path.splitext(os.path.basename(uploaded.name))[0]
            table_input = st.text_input("Destination table name", value=unique_table_name(default_name))
            c1, c2 = st.columns([1,1])
            with c1:
                do_import = st.button("Import to MySQL", type="primary")
            with c2:
                replace_existing = st.checkbox("Replace if table exists", value=False)
            if do_import:
                tn = normalize_table_name(table_input) if replace_existing else unique_table_name(table_input)
                try:
                    write_df(df, tn, replace=replace_existing)
                    st.success(f"Imported to table {tn}")
                    if st.session_state.auto_sync:
                        ok, res = rag_ingest_tables(tn)
                        if ok:
                            st.success(f"RAG synced table {tn}")
                        else:
                            st.warning(f"RAG sync failed: {res}")
                except Exception as e:
                    st.error(str(e))

with tabs[1]:
    table_list = list_tables()
    if not table_list:
        st.info("No tables found")
    else:
        left, right = st.columns([1,3])
        with left:
            selected = st.selectbox("Tables", options=table_list)
            limit = st.number_input("Limit", min_value=1, max_value=10000, value=100, step=50)
            refresh = st.button("Refresh")
        if selected:
            try:
                with engine.begin() as conn:
                    df = pd.read_sql(text(f"SELECT * FROM `{selected}` LIMIT :lim"), conn, params={"lim": int(limit)})
                st.subheader(selected)
                st.dataframe(df, use_container_width=True)
                st.download_button("Download CSV", df.to_csv(index=False).encode("utf-8"), file_name=f"{selected}.csv", mime="text/csv")
            except Exception as e:
                st.error(str(e))

with tabs[2]:
    up = st.file_uploader("Upload PDF or DOCX", type=["pdf", "docx"])
    if up:
        name = up.name.lower()
        col1, col2 = st.columns([1,1])
        with col1:
            if st.button("Save File", type="primary"):
                try:
                    up.seek(0)
                    save_file_to_db(up)
                    st.success("Saved")
                    if st.session_state.auto_sync:
                        ok, res = rag_ingest_files()
                        if ok:
                            st.success("RAG synced files")
                        else:
                            st.warning(f"RAG sync failed: {res}")
                except Exception as e:
                    st.error(str(e))
        with col2:
            if name.endswith(".pdf"):
                up.seek(0)
                b = up.read()
                b64 = base64.b64encode(b).decode("utf-8")
                components.html(f'<iframe src="data:application/pdf;base64,{b64}" width="100%" height="700px"></iframe>', height=720)
            elif name.endswith(".docx"):
                up.seek(0)
                b = up.read()
                doc = Document(io.BytesIO(b))
                text_content = "\n".join(p.text for p in doc.paragraphs)
                st.text_area("Preview", value=text_content, height=500)

with tabs[3]:
    rows = list_files()
    if not rows:
        st.info("No files stored")
    else:
        df = pd.DataFrame(rows, columns=["id","filename","content_type","size_bytes","created_at"])
        st.dataframe(df, use_container_width=True, height=300)
        sel = st.selectbox("Select file id", options=df["id"].tolist())
        if st.button("Open"):
            row = get_file_blob(sel)
            if row:
                fid, fname, ctype, sizeb, data = row
                st.write(f"{fname} • {ctype} • {sizeb} bytes")
                if ctype.startswith("application/pdf") or fname.lower().endswith(".pdf"):
                    b64 = base64.b64encode(data).decode("utf-8")
                    components.html(f'<iframe src="data:application/pdf;base64,{b64}" width="100%" height="700px"></iframe>', height=720)
                elif ctype in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",) or fname.lower().endswith(".docx"):
                    doc = Document(io.BytesIO(data))
                    text_content = "\n".join(p.text for p in doc.paragraphs)
                    st.text_area("Preview", value=text_content, height=500)
                st.download_button("Download", data, file_name=fname)

with tabs[4]:
    st.warning("This will permanently delete all tables.")
    if st.button("Clear Everything", type="primary"):
        try:
            with engine.begin() as conn:
                tables = list_tables()
                for t in tables:
                    conn.execute(text(f"DROP TABLE IF EXISTS `{t}`"))
            st.success("All tables have been cleared.")
            ensure_files_table()
            if st.session_state.auto_sync:
                ok1, res1 = rag_ingest_tables()
                ok2, res2 = rag_ingest_files()
                if ok1 and ok2:
                    st.success("RAG cleared and resynced")
                else:
                    st.warning("RAG resync encountered an issue")
        except Exception as e:
            st.error(str(e))

with tabs[5]:
    st.subheader("Chat")
    c1, c2, c3 = st.columns([2,2,1])
    with c1:
        st.text_input("LLM Model", value=os.environ.get("OLLAMA_LLM_MODEL", "mistral"), key="chat_model")
    with c2:
        st.text_input("RAG Base URL", value=rag_base, key="chat_base")
    with c3:
        if st.button("Health"):
            try:
                r = requests.get(f"{st.session_state.chat_base}/health", timeout=10)
                st.success(str(r.json()))
            except Exception as e:
                st.error(str(e))

    msg = st.text_area("Message", height=140, key="chat_msg")
    dbg = st.checkbox("Show debug details", value=True, key="chat_debug")

    if st.button("Ask"):
        if not msg.strip():
            st.warning("Enter a message")
        else:
            try:
                import time
                t0 = time.time()
                r = requests.post(
                    f"{st.session_state.chat_base}/chat",
                    json={"message": msg, "model": st.session_state.chat_model},
                    timeout=180,
                )
                dt = (time.time() - t0) * 1000.0
                if r.status_code >= 400:
                    st.error(f"Error {r.status_code}")
                    try:
                        st.code(r.text, language="json")
                    except Exception:
                        st.write(r.text)
                else:
                    data = r.json()
                    st.markdown("### Answer")
                    st.write(data.get("answer", ""))
                if dbg:
                    with st.expander("Details"):
                        st.write({"status_code": r.status_code, "latency_ms": round(dt, 1)})
                        try:
                            st.markdown("**Request Body**")
                            st.code({"message": msg, "model": st.session_state.chat_model})
                            st.markdown("**Raw Response**")
                            st.code(r.text)
                        except Exception:
                            st.write("debug display issue")
            except Exception as e:
                st.error(str(e))

